"""
SDD.md → SDD.docx 변환 스크립트
python-docx를 사용해 소프트웨어 아키텍처 명세서 Word 파일 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 설정 ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 헬퍼 함수 ──────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=10, color=None, name="맑은 고딕"):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "맑은 고딕"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 4:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_para(text="", bold=False, italic=False, size=10, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size, color=color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=9.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    return p

def add_code(text):
    """코드 블록 (회색 배경, 고정폭 폰트)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # 배경색 설정
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),  "single")
        border.set(qn("w:sz"),   "4")
        border.set(qn("w:space"),"0")
        border.set(qn("w:color"),"BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=9, color=(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, "2E74B5")
        set_cell_border(cell)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9)
            shade_cell(cell, fill)
            set_cell_border(cell)

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "4")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"BFBFBF")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("RAG Playground")
set_font(tr, bold=True, size=28, color=(0x1F, 0x49, 0x7D))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("소프트웨어 아키텍처 명세서")
set_font(sr, bold=True, size=20, color=(0x2E, 0x74, 0xB5))

doc.add_paragraph()
doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(meta_p.add_run("문서 버전: 1.0.0     |     작성일: 2026-04-06"), size=10, color=(0x60,0x60,0x60))

proj_p = doc.add_paragraph()
proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(proj_p.add_run("프로젝트명: RAG Playground — A2A 멀티 에이전트 채팅 시스템"), size=10, color=(0x60,0x60,0x60))

doc.add_paragraph()
add_divider()

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 목차
# ══════════════════════════════════════════════════════════════════════════════

add_heading("목차", level=1)

toc_items = [
    ("1.", "프로젝트 개요"),
    ("2.", "시스템 아키텍처"),
    ("3.", "기술 스택"),
    ("4.", "컴포넌트 상세 설계"),
    ("   4.1", "오케스트레이터 백엔드 (Port 8000)"),
    ("   4.2", "도메인 전문 에이전트"),
    ("   4.3", "MCP 서버"),
    ("   4.4", "프론트엔드 (Port 5173)"),
    ("5.", "데이터 모델"),
    ("6.", "통신 프로토콜 (A2A / JSON-RPC 2.0)"),
    ("7.", "API 명세"),
    ("8.", "에이전트 도구 명세"),
    ("9.", "데이터베이스 설계"),
    ("10.", "상태 흐름 및 시퀀스"),
    ("11.", "에러 처리"),
    ("12.", "환경 설정"),
    ("13.", "보안 설계"),
    ("14.", "확장성 및 운영"),
    ("15.", "비기능 요구사항"),
]
for num, text in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{num}  ")
    set_font(r1, bold=True, size=10, color=(0x2E,0x74,0xB5))
    r2 = p.add_run(text)
    set_font(r2, size=10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 프로젝트 개요
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. 프로젝트 개요", level=1)

add_heading("1.1 목적", level=2)
add_para(
    "RAG Playground는 Agent-to-Agent (A2A) 프로토콜을 기반으로 하는 멀티 에이전트 오케스트레이션 "
    "채팅 시스템이다. 단일 LLM 호출의 한계를 극복하기 위해 마스터 오케스트레이터가 도메인별 전문 "
    "에이전트에게 작업을 위임하는 계층적 위임 구조를 채택한다."
)

add_heading("1.2 핵심 기능", level=2)
add_table(
    ["기능", "설명"],
    [
        ["멀티 에이전트 오케스트레이션", "오케스트레이터가 사용자 의도를 파악해 적절한 전문 에이전트에게 위임"],
        ["실시간 스트리밍 응답", "SSE(Server-Sent Events)를 통한 토큰 단위 스트리밍"],
        ["대화 이력 관리", "MariaDB에 채팅방 및 메시지를 영속적으로 저장"],
        ["파일 처리", "PDF, DOCX, XLSX, PPTX 파일 텍스트 추출 및 문서 생성"],
        ["동적 에이전트 등록", "런타임에 새 에이전트를 레지스트리에 등록/해제"],
        ["도구 통합", "웹 검색, 날씨, 계산기, 문서 생성 등 외부 도구 사용"],
    ],
    col_widths=[5.5, 10.5]
)

add_heading("1.3 시스템 범위", level=2)
add_code(
    "사용자 (브라우저)\n"
    "    ↕ HTTP/SSE\n"
    "프론트엔드 (React, Port 5173)\n"
    "    ↕ HTTP/SSE\n"
    "오케스트레이터 (FastAPI, Port 8000)\n"
    "    ├─ MariaDB (Port 3306)\n"
    "    ├─ MCP 서버들 (Ports 8100~8102)\n"
    "    └─ 전문 에이전트들 (Ports 9000~9003)"
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. 시스템 아키텍처
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. 시스템 아키텍처", level=1)

add_heading("2.1 계층 구조 정의", level=2)
add_table(
    ["계층", "컴포넌트", "상태", "역할"],
    [
        ["프레젠테이션",   "React SPA",           "-",          "사용자 인터페이스, SSE 파싱"],
        ["오케스트레이션", "FastAPI 오케스트레이터", "Stateful",   "요청 조율, 이력 관리, 에이전트 위임"],
        ["도메인 처리",    "전문 에이전트들",        "Stateless",  "특화 도메인 쿼리 처리"],
        ["도구 레이어",    "MCP 서버들",            "Stateless",  "파일 처리, 날씨, 문서 생성"],
        ["영속성",         "MariaDB",              "-",          "채팅방 및 메시지 저장"],
    ],
    col_widths=[3.5, 4.5, 3, 5]
)

add_heading("2.2 서비스 포트 맵", level=2)
add_table(
    ["서비스", "Port", "유형"],
    [
        ["프론트엔드 (Vite 개발 서버)", "5173", "웹 UI"],
        ["오케스트레이터 (FastAPI)",    "8000", "Stateful 백엔드"],
        ["Tika MCP 서버",             "8100", "파일 처리 도구"],
        ["Weather MCP 서버",          "8101", "날씨 API 래퍼"],
        ["DocGen MCP 서버",           "8102", "문서 생성 도구"],
        ["운세 에이전트",              "9000", "도메인 에이전트 (선택)"],
        ["메이플스토리 에이전트",       "9001", "도메인 에이전트"],
        ["서든어택 에이전트",          "9002", "도메인 에이전트"],
        ["FC Online 에이전트",        "9003", "도메인 에이전트"],
        ["MariaDB",                  "3306", "관계형 DB"],
        ["Langfuse (선택)",           "3000", "LLM 관찰성"],
    ],
    col_widths=[6.5, 3, 6.5]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 기술 스택
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. 기술 스택", level=1)

add_heading("3.1 백엔드", level=2)
add_table(
    ["라이브러리", "버전", "용도"],
    [
        ["FastAPI",                "0.135+", "REST API 프레임워크"],
        ["Uvicorn",                "최신",   "ASGI 서버"],
        ["LangChain",              "0.3+",  "LLM 추상화 계층"],
        ["LangGraph",              "0.2+",  "ReAct 에이전트 그래프"],
        ["langchain-openai",       "최신",   "OpenAI ChatGPT 연동"],
        ["langchain-mcp-adapters", "최신",   "MCP 클라이언트 통합"],
        ["aiomysql",               "0.2+",  "비동기 MariaDB 드라이버"],
        ["duckduckgo-search",      "6.0+",  "웹 검색"],
        ["langfuse",               "2.0+",  "LLM 호출 관찰성"],
        ["pydantic",               "2.0+",  "데이터 검증"],
        ["httpx / aiohttp",        "최신",   "비동기 HTTP 클라이언트"],
    ],
    col_widths=[5, 2.5, 8.5]
)

add_heading("3.2 프론트엔드", level=2)
add_table(
    ["라이브러리", "버전", "용도"],
    [
        ["React",           "18.3.1", "UI 프레임워크"],
        ["Vite",            "5.4.2",  "빌드 도구 및 개발 서버"],
        ["react-markdown",  "10.1.0", "Markdown 렌더링"],
        ["Vitest",          "1.6.0+", "단위/통합 테스트 프레임워크"],
    ],
    col_widths=[5, 2.5, 8.5]
)

add_heading("3.3 MCP 서버 의존성", level=2)
add_table(
    ["라이브러리", "용도"],
    [
        ["fastmcp",      "MCP 서버 프레임워크"],
        ["pypdf",        "PDF 텍스트 추출"],
        ["python-docx",  "DOCX 처리"],
        ["openpyxl",     "XLSX 처리"],
        ["python-pptx",  "PPTX 처리"],
        ["pytesseract",  "OCR (이미지 기반 PDF)"],
        ["httpx",        "날씨 API HTTP 클라이언트"],
    ],
    col_widths=[5, 11]
)

add_heading("3.4 외부 API", level=2)
add_table(
    ["API", "용도"],
    [
        ["OpenAI API (GPT)",  "LLM 추론 엔진"],
        ["Nexon OpenAPI",     "메이플스토리/FC Online 캐릭터 데이터"],
        ["DuckDuckGo Search", "웹 검색"],
        ["OpenWeatherMap",    "날씨 정보"],
    ],
    col_widths=[5, 11]
)

add_divider()
add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. 컴포넌트 상세 설계
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. 컴포넌트 상세 설계", level=1)

# 4.1
add_heading("4.1 오케스트레이터 백엔드 (Port 8000)", level=2)

add_heading("4.1.1 모듈 구조", level=3)
add_code(
    "backend/\n"
    "├── main.py               # FastAPI 앱 초기화, 라우터 등록, lifespan 훅\n"
    "├── a2a_router.py         # POST /chat — JSON-RPC 2.0 파싱\n"
    "├── task_handler.py       # SSE 스트리밍 생성기, 에이전트 실행 조율\n"
    "├── task_store.py         # 인메모리 태스크 저장소 (UUID → Task 매핑)\n"
    "├── room_store.py         # MariaDB CRUD (채팅방, 메시지)\n"
    "├── database.py           # aiomysql 커넥션 풀 관리\n"
    "├── rooms_router.py       # GET/POST/PATCH/DELETE /rooms 엔드포인트\n"
    "├── registry_router.py    # 에이전트 레지스트리 관리 REST API\n"
    "├── agent_card.py         # GET /.well-known/agent.json\n"
    "├── models.py             # 공유 Pydantic 데이터 모델\n"
    "├── upload_router.py      # POST /upload — 파일 업로드\n"
    "├── download_router.py    # GET /download/{filename}\n"
    "└── agent/\n"
    "    ├── planner.py        # LangGraph ReAct 에이전트 빌더\n"
    "    ├── registry.py       # 서브 에이전트 탐색 및 캐싱\n"
    "    └── tools/\n"
    "        ├── delegate.py       # 서브 에이전트 위임 도구\n"
    "        ├── web_search.py     # DuckDuckGo 검색 도구\n"
    "        ├── calculator.py     # AST 기반 안전 계산기\n"
    "        ├── datetime_tool.py  # 타임존 인식 날짜/시간 도구\n"
    "        └── mcp_tools.py      # MCP 클라이언트 도구 로딩"
)

add_heading("4.1.2 요청 처리 흐름", level=3)
steps = [
    "① 프론트엔드 → POST /chat (JSON-RPC 2.0 페이로드) → a2a_router.py: JSON 파싱 → TaskSendParams 검증",
    "② task_handler.py: task_store에 Task 생성 (state: submitted) → room_store에서 대화 이력 로드 → 사용자 메시지 DB 저장 → state: working",
    "③ agent/planner.py: 시스템 프롬프트 구성 (4단계 계획 지침 + 동적 에이전트 목록) + DB 이력 변환 + 현재 메시지",
    "④ LangGraph ReAct 루프 (최대 10회): LLM 추론 → 도구 호출 결정 → on_tool_start SSE 전송 → 도구 실행 → on_chat_model_stream 토큰 스트리밍",
    "⑤ 응답 완료: 누적 응답 DB 저장 → task_artifact_update (last_chunk=true) → task_status_update (completed, final=true)",
]
for s in steps:
    add_bullet(s)

add_heading("4.1.3 LangGraph 에이전트 시스템 프롬프트 구조", level=3)
add_table(
    ["순서", "항목", "내용"],
    [
        ["1", "역할 정의",       "범용 AI 어시스턴트, 전문 에이전트 오케스트레이터"],
        ["2", "4단계 계획 원칙", "분석 → 계획 → 실행 → 종합"],
        ["3", "도구 목록",       "사용 가능한 도구 및 사용 지침"],
        ["4", "동적 에이전트 정보","현재 온라인 에이전트 URL 및 설명 (요청마다 갱신)"],
        ["5", "응답 가이드라인", "출처 인용, 불확실성 인정"],
    ],
    col_widths=[1.5, 4, 10.5]
)

add_heading("4.1.4 에이전트 레지스트리", level=3)
add_para("기본 등록 노드 (KNOWN_NODES):")
add_bullet("http://127.0.0.1:9001 — 메이플스토리 전문 에이전트")
add_bullet("http://127.0.0.1:9002 — 서든어택 전문 에이전트")
add_bullet("http://127.0.0.1:9003 — FC Online 전문 에이전트")
add_para("동작 방식:")
add_bullet("GET {base_url}/.well-known/agent.json 호출 (2초 타임아웃)")
add_bullet("성공 시 _AGENT_CACHE에 저장, 실패 시 캐시된 카드 폴백")
add_bullet("런타임 동적 등록/해제 지원 (코드 변경 불필요)")

# 4.2
add_heading("4.2 도메인 전문 에이전트", level=2)

add_heading("4.2.1 공통 특성", level=3)
add_table(
    ["속성", "값"],
    [
        ["상태",         "Stateless — DB 없음, 단일 요청 컨텍스트만 사용"],
        ["프로토콜",     "A2A JSON-RPC 2.0 (오케스트레이터와 동일)"],
        ["에이전트 광고","GET /.well-known/agent.json"],
        ["세션 ID",      "수신하지 않음 (무시)"],
    ],
    col_widths=[4, 12]
)

add_heading("4.2.2 에이전트 목록", level=3)
add_table(
    ["에이전트", "Port", "도메인", "전문 도구"],
    [
        ["메이플스토리 에이전트", "9001", "넥슨 게임 데이터",  "13개 Nexon OpenAPI 래퍼"],
        ["서든어택 에이전트",    "9002", "FPS 게임 정보",    "웹 검색, 공식 API"],
        ["FC Online 에이전트",  "9003", "축구 게임 데이터",  "넥슨 FC Online API"],
        ["운세 에이전트",        "9000", "점술/운세",        "날짜/시간, 사주 계산"],
    ],
    col_widths=[4.5, 2, 4, 5.5]
)

add_heading("4.2.3 메이플스토리 에이전트 Nexon API 도구 목록", level=3)
add_table(
    ["도구명", "API 엔드포인트", "설명"],
    [
        ["get_character_ocid",       "/maplestory/v1/id",             "캐릭터 OCID 조회"],
        ["get_character_basic",      "/v1/character/basic",           "기본 정보 (레벨, 직업, 서버)"],
        ["get_character_stat",       "/v1/character/stat",            "능력치 (HP, MP, 공격력 등)"],
        ["get_character_equipment",  "/v1/character/item-equipment",  "장착 아이템 목록"],
        ["get_character_ability",    "/v1/character/ability",         "어빌리티 정보"],
        ["get_character_hyper_stat", "/v1/character/hyper-stat",      "하이퍼스탯"],
        ["get_character_propensity", "/v1/character/propensity",      "성향 정보"],
        ["get_character_hexa_matrix","/v1/character/hexamatrix",      "헥사 매트릭스"],
        ["get_character_skill",      "/v1/character/skill",           "스킬 정보"],
        ["get_character_link_skill", "/v1/character/link-skill",      "링크 스킬"],
        ["get_character_symbol",     "/v1/character/symbol-equipment","심볼 장비"],
        ["get_character_pet",        "/v1/character/pet-equipment",   "펫 장비"],
        ["get_character_cash_item",  "/v1/character/cashitem-equipment","캐시 아이템"],
    ],
    col_widths=[5, 5.5, 5.5]
)

# 4.3
add_heading("4.3 MCP 서버", level=2)

add_heading("4.3.1 Tika MCP 서버 (Port 8100) — 파일 텍스트 추출", level=3)
add_table(
    ["지원 형식", "처리 방식"],
    [
        ["PDF",                     "pypdf 텍스트 추출 → OCR 폴백 (pytesseract)"],
        ["DOCX",                    "python-docx"],
        ["XLSX",                    "openpyxl (모든 시트)"],
        ["PPTX",                    "python-pptx (모든 슬라이드)"],
        ["TXT/MD/CSV/JSON/XML/HTML", "직접 읽기"],
    ],
    col_widths=[5, 11]
)
add_para("제공 도구:")
add_bullet("extract_text_from_file(file_path: str) → str  — 텍스트 내용 반환")
add_bullet("get_file_metadata(file_path: str) → str  — MIME 타입, 작성자, 페이지 수 등")

add_heading("4.3.2 Weather MCP 서버 (Port 8101) — OpenWeatherMap 래퍼", level=3)
add_para("제공 도구:")
add_bullet("get_current_weather(city: str, units: str = 'metric') → str")
add_bullet("get_weather_forecast(city: str, days: int = 5, units: str = 'metric') → str")
add_bullet("get_weather_by_coords(lat: float, lon: float, units: str = 'metric') → str")

add_heading("4.3.3 DocGen MCP 서버 (Port 8102) — 문서 생성", level=3)
add_para("제공 도구:")
add_bullet("create_document(filename: str, content: str, format: str) → str")
add_bullet("지원 형식: DOCX, XLSX, TXT, MD")
add_bullet("Markdown 헤더(# ~ ####) 및 리스트(-, *, •) 자동 파싱")
add_bullet("/output/ 디렉토리에 저장 후 다운로드 URL 반환")

# 4.4
add_heading("4.4 프론트엔드 (Port 5173)", level=2)

add_heading("4.4.1 컴포넌트 구조", level=3)
add_code(
    "App.jsx (루트 레이아웃 + 상태 오케스트레이션)\n"
    "├── Sidebar.jsx          채팅방 목록, 방 생성/삭제\n"
    "├── ChatWindow.jsx       메시지 표시 영역\n"
    "│   └── MessageBubble.jsx  개별 메시지 (Markdown 렌더링, 다운로드 링크)\n"
    "├── InputBar.jsx         텍스트 입력 + 파일 첨부\n"
    "└── AgentPanel.jsx       에이전트 레지스트리 관리 패널"
)

add_heading("4.4.2 커스텀 훅", level=3)
add_table(
    ["훅", "상태", "주요 기능"],
    [
        ["useRooms()",         "rooms[], activeRoomId",                    "채팅방 CRUD, 자동 초기 로드"],
        ["useA2AClient()",     "messages[], streamingText, isStreaming",   "SSE 스트리밍, 메시지 누적, abort 지원"],
        ["useAgentRegistry()", "agents[], loading",                        "10초 폴링, 에이전트 추가/삭제"],
    ],
    col_widths=[4, 5.5, 6.5]
)

add_heading("4.4.3 메시지 렌더링 규칙", level=3)
add_table(
    ["조건", "렌더링 방식"],
    [
        ["일반 텍스트/Markdown",    "react-markdown으로 렌더링"],
        ["/download/ 경로 감지",    "스타일된 다운로드 버튼 삽입"],
        ["> [도구명] 블록쿼트",     "MCP 도구: 초록색, 기타: 남색 강조"],
        ["스트리밍 중",             "끝에 깜빡이는 커서 표시"],
        ["역할 구분",               "사용자: 파란 배경, 에이전트: 어두운 배경"],
    ],
    col_widths=[5, 11]
)

add_divider()
add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. 데이터 모델
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. 데이터 모델", level=1)

add_heading("5.1 핵심 Pydantic 모델 (backend/models.py)", level=2)
add_table(
    ["클래스", "필드", "타입", "설명"],
    [
        ["TextPart",               "type",       "Literal['text']",                         "파트 타입 (고정)"],
        ["TextPart",               "text",       "str",                                     "텍스트 내용"],
        ["Message",                "role",       "Literal['user','agent']",                 "발화자 역할"],
        ["Message",                "parts",      "List[TextPart]",                          "메시지 파트 목록"],
        ["TaskStatus",             "state",      "submitted|working|completed|failed|canceled","태스크 상태"],
        ["Task",                   "id",         "str (UUID)",                              "태스크 고유 ID"],
        ["Task",                   "session_id", "Optional[str]",                           "채팅방 ID (Stateful)"],
        ["Artifact",               "index",      "int",                                     "청크 순번"],
        ["Artifact",               "last_chunk", "bool",                                    "마지막 청크 여부"],
        ["JsonRpcRequest",         "jsonrpc",    "Literal['2.0']",                          "프로토콜 버전"],
        ["JsonRpcRequest",         "method",     "str",                                     "RPC 메서드명"],
        ["TaskSendParams",         "session_id", "Optional[str]",                           "채팅방 ID"],
        ["TaskSendParams",         "message",    "Message",                                 "사용자 메시지"],
        ["TaskStatusUpdateEvent",  "type",       "Literal['task_status_update']",           "이벤트 타입"],
        ["TaskStatusUpdateEvent",  "final",      "bool",                                    "스트림 종료 여부"],
        ["TaskArtifactUpdateEvent","type",       "Literal['task_artifact_update']",         "이벤트 타입"],
        ["TaskArtifactUpdateEvent","artifact",   "Artifact",                                "응답 청크 데이터"],
    ],
    col_widths=[4, 3, 4.5, 4.5]
)

add_heading("5.2 에이전트 카드 스키마", level=2)
add_code(
    '{\n'
    '  "name":        string,   // 에이전트 표시명\n'
    '  "description": string,   // 기능 설명\n'
    '  "url":         string,   // A2A 엔드포인트 URL\n'
    '  "version":     string,   // 버전\n'
    '  "provider": {\n'
    '    "organization": string,\n'
    '    "model":        string\n'
    '  },\n'
    '  "skills": [\n'
    '    { "id": string, "name": string, "description": string }\n'
    '  ]\n'
    '}'
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. 통신 프로토콜
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. 통신 프로토콜 (A2A / JSON-RPC 2.0)", level=1)

add_heading("6.1 요청 형식", level=2)
add_code(
    '{\n'
    '  "jsonrpc": "2.0",\n'
    '  "id": "550e8400-e29b-41d4-a716-446655440000",\n'
    '  "method": "message/stream",\n'
    '  "params": {\n'
    '    "session_id": "room-uuid",\n'
    '    "message": {\n'
    '      "role": "user",\n'
    '      "parts": [{"type": "text", "text": "메이플스토리 캐릭터 조회"}]\n'
    '    },\n'
    '    "metadata": { "model": "gpt-5.4-mini" }\n'
    '  }\n'
    '}'
)
add_para("※ 서브 에이전트 위임 시 session_id 필드 생략 (Stateless)", italic=True, color=(0x60,0x60,0x60))

add_heading("6.2 SSE 응답 이벤트 타입", level=2)
add_table(
    ["이벤트 타입", "final", "용도"],
    [
        ["task_status_update",   "false", "상태 전환 알림 (submitted → working)"],
        ["task_status_update",   "true",  "스트리밍 완료 신호 (completed)"],
        ["task_artifact_update", "—",     "LLM 토큰 / 도구 호출 공지 / 다운로드 링크"],
        ["error",                "—",     "처리 오류 발생 시 전송"],
    ],
    col_widths=[5, 2, 9]
)
add_para("HTTP 응답 헤더:")
add_bullet("Content-Type: text/event-stream")
add_bullet("Cache-Control: no-cache")
add_bullet("X-Accel-Buffering: no")

add_heading("6.3 지원 메서드", level=2)
add_table(
    ["메서드", "설명"],
    [["message/stream", "SSE 스트리밍 채팅 요청 (유일하게 지원)"]],
    col_widths=[5, 11]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. API 명세
# ══════════════════════════════════════════════════════════════════════════════

add_heading("7. API 명세", level=1)

add_heading("7.1 A2A 및 에이전트 카드 엔드포인트", level=2)
add_table(
    ["메서드", "경로", "설명"],
    [
        ["POST", "/chat",                      "A2A JSON-RPC 2.0 채팅 (SSE 스트리밍)"],
        ["GET",  "/.well-known/agent.json",    "에이전트 카드 광고"],
    ],
    col_widths=[2, 5, 9]
)

add_heading("7.2 채팅방 API", level=2)
add_table(
    ["메서드", "경로", "설명", "요청 바디"],
    [
        ["GET",    "/rooms",                    "채팅방 목록 (updated_at DESC)", "—"],
        ["POST",   "/rooms",                    "새 채팅방 생성",                '{"title": str, "model": str}'],
        ["GET",    "/rooms/{room_id}",          "채팅방 상세 조회",              "—"],
        ["PATCH",  "/rooms/{room_id}",          "채팅방 제목 수정",              '{"title": str}'],
        ["DELETE", "/rooms/{room_id}",          "채팅방 삭제 (메시지 CASCADE)",  "—"],
        ["GET",    "/rooms/{room_id}/messages", "메시지 이력 (created_at ASC)",  "—"],
    ],
    col_widths=[2, 5, 4.5, 4.5]
)

add_heading("7.3 에이전트 레지스트리 API", level=2)
add_table(
    ["메서드", "경로", "설명", "요청 바디"],
    [
        ["GET",    "/registry/agents", "에이전트 목록 + 온라인 상태",  "—"],
        ["POST",   "/registry/agents", "새 에이전트 URL 등록",         '{"url": str}'],
        ["DELETE", "/registry/agents", "에이전트 URL 제거",            '{"url": str}'],
    ],
    col_widths=[2, 5, 5, 4]
)

add_heading("7.4 파일 API", level=2)
add_table(
    ["메서드", "경로", "설명"],
    [
        ["POST", "/upload",              "파일 업로드 (multipart/form-data, 최대 50MB)"],
        ["GET",  "/download/{filename}", "파일 다운로드"],
    ],
    col_widths=[2, 5, 9]
)
add_para("허용 확장자: .pdf .docx .xlsx .pptx .txt .md .csv .json .xml .html")

add_divider()
add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. 에이전트 도구 명세
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8. 에이전트 도구 명세", level=1)

add_heading("8.1 오케스트레이터 기본 도구", level=2)
add_table(
    ["도구명", "시그니처", "주요 특성"],
    [
        ["get_datetime",   "get_datetime(timezone='Asia/Seoul') → str",  "ZoneInfo 기반, 한국어 요일 포함"],
        ["calculator",     "calculator(expression: str) → str",          "AST 파싱 (eval 미사용), ^ 별칭 지원"],
        ["web_search",     "web_search(query: str) → str",               "DuckDuckGo, kr-kr 지역, 최대 4건"],
        ["delegate_task",  "delegate_task(target_url, task_description) → str", "A2A SSE 파싱, 15초 타임아웃"],
    ],
    col_widths=[3.5, 6, 6.5]
)

add_heading("8.2 MCP 도구", level=2)
add_table(
    ["도구명", "MCP 서버", "설명"],
    [
        ["extract_text_from_file", "Tika (8100)",    "파일 경로 → 텍스트 내용"],
        ["get_file_metadata",      "Tika (8100)",    "파일 경로 → 메타데이터"],
        ["get_current_weather",    "Weather (8101)", "도시명 → 현재 날씨"],
        ["get_weather_forecast",   "Weather (8101)", "도시명, 일수 → 예보"],
        ["get_weather_by_coords",  "Weather (8101)", "위경도 → 현재 날씨"],
        ["create_document",        "DocGen (8102)",  "파일명, 내용, 형식 → 다운로드 URL"],
    ],
    col_widths=[5, 4, 7]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 9. 데이터베이스 설계
# ══════════════════════════════════════════════════════════════════════════════

add_heading("9. 데이터베이스 설계", level=1)
add_para("적용 대상: 오케스트레이터만 (서브 에이전트는 DB 없음)")

add_heading("9.1 rooms 테이블", level=2)
add_table(
    ["컬럼", "타입", "제약", "설명"],
    [
        ["id",         "CHAR(36)",    "PRIMARY KEY",          "UUID"],
        ["title",      "VARCHAR(100)","DEFAULT 'New Chat'",   "채팅방 제목 (첫 메시지 40자 자동 설정)"],
        ["model",      "VARCHAR(50)", "DEFAULT 'gpt-5.4-mini'","사용 모델"],
        ["created_at", "BIGINT",      "NOT NULL",             "생성 시각 (Unix 밀리초)"],
        ["updated_at", "BIGINT",      "NOT NULL, INDEX DESC", "최근 활동 시각 (메시지마다 갱신)"],
    ],
    col_widths=[3, 3, 4.5, 5.5]
)

add_heading("9.2 messages 테이블", level=2)
add_table(
    ["컬럼", "타입", "제약", "설명"],
    [
        ["id",         "CHAR(36)",           "PRIMARY KEY",         "UUID"],
        ["room_id",    "CHAR(36)",           "FK → rooms(id) CASCADE","채팅방 참조"],
        ["role",       "ENUM('user','agent')","NOT NULL",            "발화자 역할"],
        ["content",    "TEXT",               "NOT NULL",            "메시지 전체 텍스트"],
        ["created_at", "BIGINT",             "NOT NULL, INDEX ASC", "전송 시각 (이력 조회용)"],
    ],
    col_widths=[3, 4, 4, 5]
)

add_heading("9.3 커넥션 풀 설정", level=2)
add_table(
    ["설정값", "값"],
    [
        ["minsize",    "2"],
        ["maxsize",    "10"],
        ["charset",    "utf8mb4"],
        ["autocommit", "True"],
    ],
    col_widths=[5, 11]
)

add_heading("9.4 주요 DB 연산", level=2)
add_table(
    ["함수", "SQL", "설명"],
    [
        ["create_room()",       "INSERT INTO rooms",                                 "UUID 생성 후 삽입"],
        ["list_rooms()",        "SELECT * ORDER BY updated_at DESC",                 "최근 활동순 목록"],
        ["get_messages(id)",    "SELECT * WHERE room_id ORDER BY created_at ASC",    "대화 이력 조회"],
        ["save_message(...)",   "INSERT INTO messages + UPDATE rooms SET updated_at","메시지 저장 + 방 갱신"],
        ["update_room_title()", "UPDATE rooms SET title",                            "첫 메시지 자동 제목"],
        ["delete_room(id)",     "DELETE FROM rooms",                                 "메시지 CASCADE 삭제"],
    ],
    col_widths=[4, 6, 6]
)

add_divider()
add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. 상태 흐름 및 시퀀스
# ══════════════════════════════════════════════════════════════════════════════

add_heading("10. 상태 흐름 및 시퀀스", level=1)

add_heading("10.1 태스크 상태 전이", level=2)
add_code(
    "[submitted] → [working] → [completed]\n"
    "                       ↘ [failed]\n"
    "           ↘ [canceled]  (abort 신호 수신 시)"
)

add_heading("10.2 전체 채팅 시퀀스 (서브 에이전트 위임 포함)", level=2)
add_code(
    "사용자     프론트엔드    오케스트레이터   서브에이전트    MariaDB\n"
    "  │            │               │              │             │\n"
    "  │ 메시지입력  │               │              │             │\n"
    "  │───────────>│               │              │             │\n"
    "  │            │  POST /chat   │              │             │\n"
    "  │            │──────────────>│              │             │\n"
    "  │            │               │  이력 로드   │             │\n"
    "  │            │               │─────────────────────────────>\n"
    "  │            │               │<────────────────────────────\n"
    "  │            │  SSE:submitted│              │             │\n"
    "  │            │<──────────────│              │             │\n"
    "  │            │  SSE:working  │              │             │\n"
    "  │            │<──────────────│              │             │\n"
    "  │            │               │ 사용자메시지저장             │\n"
    "  │            │               │─────────────────────────────>\n"
    "  │            │               │  LLM 추론    │             │\n"
    "  │            │  SSE:도구알림 │              │             │\n"
    "  │            │<──────────────│              │             │\n"
    "  │            │               │  delegate_task             │\n"
    "  │            │               │─────────────>│             │\n"
    "  │            │               │<─────────────│             │\n"
    "  │            │  SSE:토큰들   │              │             │\n"
    "  │            │<──────────────│              │             │\n"
    "  │ UI실시간출력│               │ 에이전트응답저장             │\n"
    "  │<───────────│               │─────────────────────────────>\n"
    "  │            │  SSE:completed│              │             │\n"
    "  │            │<──────────────│              │             │"
)

add_heading("10.3 에이전트 레지스트리 갱신 흐름", level=2)
add_code(
    "프론트엔드                  오케스트레이터              서브 에이전트들\n"
    "    │                           │                          │\n"
    "    │  [10초마다]                │                          │\n"
    "    │  GET /registry/agents      │                          │\n"
    "    │──────────────────────────>│                          │\n"
    "    │                           │  GET /.well-known/agent.json\n"
    "    │                           │─────────────────────────>│ (병렬, 2초)\n"
    "    │                           │<─────────────────────────│\n"
    "    │  [{name, url, online,...}]│                          │\n"
    "    │<──────────────────────────│                          │\n"
    "    │  UI 상태 갱신              │                          │"
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 11. 에러 처리
# ══════════════════════════════════════════════════════════════════════════════

add_heading("11. 에러 처리", level=1)

add_heading("11.1 JSON-RPC 에러 코드", level=2)
add_table(
    ["코드", "의미", "발생 위치"],
    [
        ["-32700", "Parse error (JSON 파싱 실패)",            "a2a_router.py"],
        ["-32601", "Method not found (message/stream 외)",   "a2a_router.py"],
        ["-32602", "Invalid params (TaskSendParams 검증 실패)","a2a_router.py"],
        ["-32603", "Internal server error (LangGraph 예외 등)","task_handler.py"],
    ],
    col_widths=[2.5, 7, 6.5]
)

add_heading("11.2 특수 에러 처리", level=2)
add_table(
    ["상황", "처리 방식"],
    [
        ["GraphRecursionError (도구 10회 초과)", "사용자 친화적 메시지: '도구 호출 한도에 도달했습니다'"],
        ["서브 에이전트 타임아웃 (15초)",         "에러 문자열을 LLM에 반환, LLM이 사용자에게 안내"],
        ["파일 업로드 크기 초과",               "HTTP 413 응답"],
        ["허용되지 않는 파일 형식",             "HTTP 400 응답"],
        ["에이전트 카드 조회 실패",             "캐시된 카드 폴백, 없으면 목록 제외"],
        ["프론트엔드 AbortError",              "조용히 무시 (사용자 직접 중단)"],
    ],
    col_widths=[5.5, 10.5]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 12. 환경 설정
# ══════════════════════════════════════════════════════════════════════════════

add_heading("12. 환경 설정", level=1)

add_heading("12.1 .env 파일 항목", level=2)
add_table(
    ["항목", "설명"],
    [
        ["OPENAI_API_KEY",        "OpenAI API 키"],
        ["MODEL_NAME",            "사용 모델명 (예: gpt-5.4-mini)"],
        ["DB_HOST / DB_PORT",     "MariaDB 호스트/포트"],
        ["DB_NAME / DB_USER / DB_PASSWORD", "MariaDB 데이터베이스/계정"],
        ["LANGFUSE_PUBLIC_KEY / LANGFUSE_SECRET_KEY", "Langfuse LLM 관찰성 키"],
        ["LANGFUSE_HOST",         "Langfuse 서버 URL"],
        ["OPENWEATHER_API_KEY",   "OpenWeatherMap API 키"],
        ["NEXON_API_KEY",         "Nexon OpenAPI 키"],
        ["TIKA_MCP_URL",          "Tika MCP 서버 URL (http://127.0.0.1:8100/sse)"],
        ["WEATHER_MCP_URL",       "Weather MCP 서버 URL (http://127.0.0.1:8101/sse)"],
        ["DOCGEN_MCP_URL",        "DocGen MCP 서버 URL (http://127.0.0.1:8102/sse)"],
        ["BACKEND_URL",           "DocGen 서버용 백엔드 URL (다운로드 링크 생성)"],
    ],
    col_widths=[6, 10]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 13. 보안 설계
# ══════════════════════════════════════════════════════════════════════════════

add_heading("13. 보안 설계", level=1)

add_heading("13.1 적용된 보안 조치", level=2)
add_table(
    ["항목", "조치"],
    [
        ["CORS",         "localhost:5173만 허용 (프로덕션 배포 시 수정 필요)"],
        ["계산기 안전성", "eval() 미사용, AST 파싱 기반 수식 평가"],
        ["파일 업로드",   "확장자 화이트리스트, 50MB 크기 제한"],
        ["SQL 인젝션",    "aiomysql 파라미터 바인딩 (%s 플레이스홀더)"],
        ["API 키 보안",   ".env 파일로 분리, 코드에 하드코딩 없음"],
        ["에이전트 검증", "/.well-known/agent.json 응답 스키마 검증"],
    ],
    col_widths=[4, 12]
)

add_heading("13.2 향후 보안 개선 권고사항", level=2)
add_table(
    ["항목", "현황", "권고"],
    [
        ["인증/인가",      "미적용",          "JWT 또는 세션 기반 인증 추가"],
        ["요청 레이트 리밋","미적용",          "IP/사용자 기반 제한"],
        ["HTTPS",         "미적용 (개발 환경)", "프로덕션: TLS 종료 (Nginx/Traefik)"],
        ["에이전트 신뢰검증","미적용",          "HMAC 서명 또는 mTLS"],
    ],
    col_widths=[3.5, 4, 8.5]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 14. 확장성 및 운영
# ══════════════════════════════════════════════════════════════════════════════

add_heading("14. 확장성 및 운영", level=1)

add_heading("14.1 새 에이전트 추가 절차", level=2)
steps = [
    "maplestory-agent/ 구조를 복사",
    "agent_card.py에서 이름, 설명, 포트 변경",
    "agent/planner.py에서 도메인 전용 시스템 프롬프트 작성",
    "agent/tools/에 전문 도구 추가",
    "main.py에서 포트 설정",
    "서버 실행 후 오케스트레이터의 /registry/agents에 URL 등록 (코드 변경 불필요)",
]
for i, s in enumerate(steps, 1):
    add_bullet(f"Step {i}. {s}")

add_heading("14.2 새 MCP 서버 추가 절차", level=2)
steps2 = [
    "fastmcp 기반 서버 작성 (server.py)",
    "@mcp.tool() 데코레이터로 도구 정의",
    ".env에 NEW_MCP_URL=http://127.0.0.1:81xx/sse 추가",
    "backend/agent/tools/mcp_tools.py의 MCP_SERVER_URLS에 URL 추가",
    "시스템 프롬프트에 새 도구 설명 추가",
]
for i, s in enumerate(steps2, 1):
    add_bullet(f"Step {i}. {s}")

add_heading("14.3 서비스별 실행 명령", level=2)
add_code(
    "# 데이터베이스 초기화\n"
    "mariadb -u root -p < schema.sql\n\n"
    "# MCP 서버들\n"
    "cd tika-mcp-server    && python server.py &\n"
    "cd weather-mcp-server && python server.py &\n"
    "cd docgen-mcp-server  && python server.py &\n\n"
    "# 서브 에이전트들\n"
    "cd maplestory-agent   && python main.py &\n"
    "cd suddenattack-agent && python main.py &\n"
    "cd fconline-agent     && python main.py &\n\n"
    "# 오케스트레이터\n"
    "cd backend && python main.py\n\n"
    "# 프론트엔드 (개발)\n"
    "cd frontend && npm install && npm run dev"
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 15. 비기능 요구사항
# ══════════════════════════════════════════════════════════════════════════════

add_heading("15. 비기능 요구사항", level=1)

add_heading("15.1 성능 목표", level=2)
add_table(
    ["항목", "목표"],
    [
        ["첫 토큰 응답 시간 (TTFT)", "< 2초 (LLM API 포함)"],
        ["서브 에이전트 위임 타임아웃", "15초"],
        ["에이전트 카드 조회 타임아웃", "2초"],
        ["DB 커넥션 풀",              "min 2, max 10"],
        ["LangGraph 최대 반복 횟수",  "10회 (GraphRecursionError 방지)"],
    ],
    col_widths=[6, 10]
)

add_heading("15.2 가용성", level=2)
add_table(
    ["항목", "현황"],
    [
        ["에이전트 오프라인 폴백",  "캐시된 카드 사용, 오케스트레이터 계속 동작"],
        ["MCP 서버 장애",          "해당 도구 미제공으로 처리, 시스템 계속 동작"],
        ["DB 커넥션 오류",         "에러 응답 반환 (graceful degradation)"],
    ],
    col_widths=[6, 10]
)

add_heading("15.3 관찰성", level=2)
add_table(
    ["항목", "도구"],
    [
        ["LLM 호출 추적",   "Langfuse (호출 횟수, 토큰 사용량, 레이턴시)"],
        ["에러 로깅",       "콘솔 로깅 (uvicorn access log + Python logging)"],
        ["에이전트 상태",   "프론트엔드 AgentPanel (10초 폴링)"],
    ],
    col_widths=[5, 11]
)

add_heading("15.4 유지보수성", level=2)
add_para("• 모듈화: 각 에이전트/MCP 서버는 독립적으로 배포/업데이트 가능")
add_para("• 동적 등록: 코드 변경 없이 에이전트 추가/제거 가능")
add_para("• 공유 모델: models.py로 데이터 모델 중앙 관리")
add_para("• 설정 분리: 모든 민감 설정은 .env에 격리")

add_divider()

# ── 꼬리말
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer_p.add_run("본 문서는 RAG Playground v1.0.0 기준으로 작성되었습니다."),
         italic=True, size=9, color=(0x80, 0x80, 0x80))

# ── 저장 ───────────────────────────────────────────────────────────────────────
output_path = "C:/sjkim/rag-playground/docs/SDD_아키텍처명세서.docx"
doc.save(output_path)
print(f"[완료] {output_path}")
