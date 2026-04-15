"""
DocGen MCP Server — 문서 생성 및 다운로드 MCP 서버.

지원 형식:
  - TXT  : 평문 텍스트
  - MD   : 마크다운
  - DOCX : Word 문서 (python-docx)
  - XLSX : Excel 스프레드시트 (openpyxl)

실행:
  python server.py
  → SSE 서버: http://0.0.0.0:8102/sse
"""
import os
import re
import logging
import uuid
from pathlib import Path

from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP

load_dotenv(Path(__file__).parent.parent / ".env")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MCP_HOST = os.getenv("MCP_HOST", "0.0.0.0")
MCP_PORT = int(os.getenv("MCP_PORT", "8102"))
BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8000")

OUTPUT_DIR = Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

mcp = FastMCP("DocGen", host=MCP_HOST, port=MCP_PORT)


def _safe_stem(name: str) -> str:
    """파일명에서 위험 문자 제거."""
    stem = Path(name).stem if "." in name else name
    stem = re.sub(r'[\\/*?:"<>|]', "_", stem).strip(". ")
    return stem or "document"


def _unique_path(stem: str, suffix: str) -> Path:
    filename = f"{stem}{suffix}"
    path = OUTPUT_DIR / filename
    if path.exists():
        filename = f"{stem}_{uuid.uuid4().hex[:6]}{suffix}"
        path = OUTPUT_DIR / filename
    return path


def _download_url(filename: str) -> str:
    return f"{BACKEND_URL}/download/{filename}"


# ── DOCX 헬퍼 ────────────────────────────────────────────────────────────────

def _build_docx(content: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 기본 스타일 조정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # 제목 감지
        h_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)
            doc.add_heading(heading_text, level=min(level, 4))
            continue

        # 구분선
        if re.match(r"^[-=]{3,}$", stripped):
            doc.add_paragraph("─" * 40)
            continue

        # 번호 목록
        ol_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if ol_match:
            p = doc.add_paragraph(ol_match.group(1), style="List Number")
            continue

        # 불릿 목록
        ul_match = re.match(r"^[-*•]\s+(.*)", stripped)
        if ul_match:
            p = doc.add_paragraph(ul_match.group(1), style="List Bullet")
            continue

        # 일반 문단 — 인라인 볼드/이탤릭 처리
        p = doc.add_paragraph()
        _add_inline(p, stripped)

    doc.save(str(path))


def _add_inline(paragraph, text: str) -> None:
    """**bold** 와 *italic* 을 Run으로 분리해서 추가."""
    from docx.shared import Pt
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(0).startswith("**"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        else:
            run = paragraph.add_run(m.group(3))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── XLSX 헬퍼 ────────────────────────────────────────────────────────────────

def _build_xlsx(content: str, path: Path) -> None:
    """CSV 또는 마크다운 테이블 형식의 content를 XLSX로 변환."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active

    rows = []
    for line in content.splitlines():
        line = line.strip()
        if not line or re.match(r"^[\|\s\-:]+$", line):
            continue
        # 마크다운 테이블 또는 CSV 파싱
        if "|" in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
        else:
            cells = [c.strip() for c in line.split(",")]
        rows.append(cells)

    header_fill = PatternFill("solid", fgColor="2563EB")
    header_font = Font(bold=True, color="FFFFFF")

    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if r_idx == 1:
                cell.font = header_font
                cell.fill = header_fill

    # 열 너비 자동 조정
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    wb.save(str(path))


# ── MCP 툴 ───────────────────────────────────────────────────────────────────

@mcp.tool()
def create_document(content: str, filename: str, format: str = "docx") -> str:
    """
    텍스트 내용으로 문서 파일을 생성하고 다운로드 링크를 반환합니다.

    Args:
        content : 문서에 들어갈 내용. 마크다운 형식 지원 (제목 #, 목록 -, 볼드 **).
                  Excel의 경우 CSV 또는 마크다운 테이블 형식.
        filename: 저장할 파일명 (확장자 없이). 예: "월간보고서", "회의록"
        format  : 파일 형식. "txt" / "md" / "docx" (기본값) / "xlsx"

    Returns:
        생성된 파일의 다운로드 URL과 안내 메시지.
    """
    fmt = format.lower().lstrip(".")
    ext_map = {"txt": ".txt", "md": ".md", "docx": ".docx", "xlsx": ".xlsx"}
    if fmt not in ext_map:
        return f"지원하지 않는 형식입니다: {format}. 지원 형식: txt, md, docx, xlsx"

    suffix = ext_map[fmt]
    stem = _safe_stem(filename)
    path = _unique_path(stem, suffix)

    try:
        if fmt in ("txt", "md"):
            path.write_text(content, encoding="utf-8")
        elif fmt == "docx":
            _build_docx(content, path)
        elif fmt == "xlsx":
            _build_xlsx(content, path)
    except Exception as e:
        logger.error(f"문서 생성 실패: {e}")
        return f"문서 생성 중 오류가 발생했습니다: {e}"

    url = _download_url(path.name)
    logger.info(f"문서 생성 완료: {path.name} ({path.stat().st_size} bytes)")
    return (
        f"✅ 문서가 생성되었습니다.\n"
        f"- 파일명: {path.name}\n"
        f"- 형식: {suffix.upper()}\n"
        f"- 다운로드: {url}"
    )


if __name__ == "__main__":
    logger.info(f"DocGen MCP 서버 시작: http://{MCP_HOST}:{MCP_PORT}/sse")
    mcp.run(transport="sse")
